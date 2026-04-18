"""
Mentorship Module - Create complete mentorship programs from user knowledge
"""
from fastapi import APIRouter, Request, HTTPException, UploadFile, File, Form
from pydantic import BaseModel
from typing import List, Optional
import uuid, os, json
from datetime import datetime, timezone

router = APIRouter(prefix="/api/mentorship")

db = None
get_current_user = None
llm_generate = None  # Will be set from server.py

def init(database, auth_fn, llm_fn):
    global db, get_current_user, llm_generate
    db = database
    get_current_user = auth_fn
    llm_generate = llm_fn

MENTORSHIP_PROMPT = """Voce e um especialista em criacao de mentorias e infoprodutos digitais. Com base no conhecimento fornecido pelo usuario, crie uma MENTORIA COMPLETA e detalhada.

## CONHECIMENTO DO USUARIO:
{knowledge}

## INSTRUCOES:
Crie uma mentoria completa com a seguinte estrutura:

1. **NOME DA MENTORIA** - Um nome atrativo e profissional
2. **PROMESSA PRINCIPAL** - O que o aluno vai conquistar (1 frase poderosa)
3. **PUBLICO-ALVO** - Para quem e essa mentoria (perfil detalhado)
4. **DURACAO E FORMATO** - Semanas, encontros, formato (ao vivo, gravado, misto)
5. **MODULOS DETALHADOS** - Cada modulo com:
   - Nome do modulo
   - Objetivo
   - Aulas (titulo + descricao de cada aula)
   - Materiais complementares
   - Exercicio pratico
6. **BONUS** - Materiais extras, templates, checklists
7. **METODOLOGIA** - Como o aluno vai aprender (passo a passo)
8. **RESULTADOS ESPERADOS** - O que o aluno alcanca ao final de cada modulo
9. **FAQ** - Perguntas frequentes e respostas
10. **COPY DE VENDAS** - Headline, subheadline, bullets de beneficio, CTA
11. **ESTRATEGIA DE PRECO** - Sugestao de precificacao com ancoragem

Seja extremamente detalhado e profissional. Cada modulo deve ter no minimo 4-5 aulas. A mentoria deve ter no minimo 6 modulos.
Responda em portugues brasileiro."""

class MentorshipCreate(BaseModel):
    title: Optional[str] = ""
    knowledge_text: str = ""
    niche: str = ""
    target_audience: str = ""
    duration_weeks: int = 8

class MentorshipUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    status: Optional[str] = None
    modules: Optional[list] = None

class ModuleData(BaseModel):
    id: Optional[str] = None
    title: str
    objective: str = ""
    order: int = 0
    lessons: list = []  # [{id, title, content, duration, order}]
    exercises: list = []  # [{id, title, description}]
    materials: list = []  # [{id, title, type}]

# ─── Knowledge Upload ────────────────────────────────────────────────────────
@router.post("/upload-knowledge")
async def upload_knowledge(request: Request, file: UploadFile = File(...)):
    """Upload a file with user's knowledge (txt, pdf, md, doc)."""
    user = await get_current_user(request)
    content = await file.read()
    text = ""
    filename = file.filename or "upload"

    if filename.endswith(('.txt', '.md')):
        text = content.decode('utf-8', errors='ignore')
    elif filename.endswith('.csv'):
        text = content.decode('utf-8', errors='ignore')
    else:
        # Try to decode as text
        try:
            text = content.decode('utf-8', errors='ignore')
        except Exception:
            text = str(content[:5000])

    # Store knowledge
    knowledge_id = str(uuid.uuid4())
    doc = {
        "id": knowledge_id,
        "user_id": user["_id"],
        "filename": filename,
        "content": text[:50000],  # Limit to 50k chars
        "size": len(text),
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.knowledge_base.insert_one(doc)
    doc.pop("_id", None)
    return {"id": knowledge_id, "filename": filename, "size": len(text), "preview": text[:300]}

@router.get("/knowledge")
async def list_knowledge(request: Request):
    user = await get_current_user(request)
    items = await db.knowledge_base.find(
        {"user_id": user["_id"]}, {"_id": 0, "content": 0}
    ).sort("created_at", -1).to_list(20)
    return items

@router.delete("/knowledge/{kid}")
async def delete_knowledge(kid: str, request: Request):
    user = await get_current_user(request)
    await db.knowledge_base.delete_one({"id": kid, "user_id": user["_id"]})
    return {"message": "Conhecimento removido"}

# ─── Mentorship CRUD ─────────────────────────────────────────────────────────
@router.get("/list")
async def list_mentorships(request: Request):
    user = await get_current_user(request)
    mentorships = await db.mentorships.find(
        {"user_id": user["_id"]}, {"_id": 0}
    ).sort("created_at", -1).to_list(20)
    return mentorships

@router.get("/{mentorship_id}")
async def get_mentorship(mentorship_id: str, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]}, {"_id": 0})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")
    return m

@router.post("/generate")
async def generate_mentorship(body: MentorshipCreate, request: Request):
    """Generate a complete mentorship using AI based on user's knowledge."""
    user = await get_current_user(request)

    # Gather knowledge from text + uploaded files
    knowledge_parts = []
    if body.knowledge_text:
        knowledge_parts.append(body.knowledge_text)

    # Also pull from knowledge base
    kb_items = await db.knowledge_base.find({"user_id": user["_id"]}).to_list(10)
    for item in kb_items:
        knowledge_parts.append(f"[Arquivo: {item.get('filename', '')}]\n{item.get('content', '')[:10000]}")

    if not knowledge_parts:
        raise HTTPException(status_code=400, detail="Forneca seu conhecimento no campo de texto ou faca upload de um arquivo")

    full_knowledge = "\n\n---\n\n".join(knowledge_parts)
    if body.niche:
        full_knowledge += f"\n\nNicho: {body.niche}"
    if body.target_audience:
        full_knowledge += f"\nPublico-alvo desejado: {body.target_audience}"
    if body.duration_weeks:
        full_knowledge += f"\nDuracao desejada: {body.duration_weeks} semanas"

    # Generate via LLM
    prompt = MENTORSHIP_PROMPT.format(knowledge=full_knowledge[:15000])
    content = await llm_generate(prompt, user["_id"])

    # Store mentorship
    mentorship_id = str(uuid.uuid4())
    title = body.title or "Nova Mentoria"
    # Try to extract title from generated content
    if "**NOME DA MENTORIA**" in content:
        try:
            title_line = content.split("**NOME DA MENTORIA**")[1].split("\n")[0].strip().strip(":-–— ")
            if title_line:
                title = title_line
        except Exception:
            pass

    doc = {
        "id": mentorship_id,
        "user_id": user["_id"],
        "title": title,
        "content": content,
        "knowledge_summary": full_knowledge[:1000],
        "niche": body.niche,
        "target_audience": body.target_audience,
        "duration_weeks": body.duration_weeks,
        "status": "draft",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    await db.mentorships.insert_one(doc)
    doc.pop("_id", None)
    return doc

@router.put("/{mentorship_id}")
async def update_mentorship(mentorship_id: str, body: MentorshipUpdate, request: Request):
    user = await get_current_user(request)
    data = {k: v for k, v in body.model_dump().items() if v is not None}
    if data:
        await db.mentorships.update_one({"id": mentorship_id, "user_id": user["_id"]}, {"$set": data})
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]}, {"_id": 0})
    return m

@router.delete("/{mentorship_id}")
async def delete_mentorship(mentorship_id: str, request: Request):
    user = await get_current_user(request)
    await db.mentorships.delete_one({"id": mentorship_id, "user_id": user["_id"]})
    return {"message": "Mentoria deletada"}

# ─── Modules CRUD ────────────────────────────────────────────────────────────
@router.get("/{mentorship_id}/modules")
async def get_modules(mentorship_id: str, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")
    return m.get("modules", [])

@router.put("/{mentorship_id}/modules")
async def save_modules(mentorship_id: str, request: Request):
    """Save the full modules structure (replace all)."""
    user = await get_current_user(request)
    body = await request.json()
    modules = body if isinstance(body, list) else body.get("modules", [])
    # Ensure IDs
    for i, mod in enumerate(modules):
        if not mod.get("id"):
            mod["id"] = str(uuid.uuid4())
        mod["order"] = i
        for j, lesson in enumerate(mod.get("lessons", [])):
            if not lesson.get("id"):
                lesson["id"] = str(uuid.uuid4())
            lesson["order"] = j
    await db.mentorships.update_one(
        {"id": mentorship_id, "user_id": user["_id"]},
        {"$set": {"modules": modules}}
    )
    return modules

@router.post("/{mentorship_id}/modules")
async def add_module(mentorship_id: str, body: ModuleData, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")
    modules = m.get("modules", [])
    mod = body.model_dump()
    mod["id"] = str(uuid.uuid4())
    mod["order"] = len(modules)
    for j, lesson in enumerate(mod.get("lessons", [])):
        if not lesson.get("id"):
            lesson["id"] = str(uuid.uuid4())
        lesson["order"] = j
    modules.append(mod)
    await db.mentorships.update_one({"id": mentorship_id}, {"$set": {"modules": modules}})
    return mod

@router.delete("/{mentorship_id}/modules/{module_id}")
async def delete_module(mentorship_id: str, module_id: str, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")
    modules = [mod for mod in m.get("modules", []) if mod.get("id") != module_id]
    for i, mod in enumerate(modules):
        mod["order"] = i
    await db.mentorships.update_one({"id": mentorship_id}, {"$set": {"modules": modules}})
    return {"message": "Modulo removido"}

# ─── Parse AI content into structured modules ────────────────────────────────
@router.post("/{mentorship_id}/parse-modules")
async def parse_content_to_modules(mentorship_id: str, request: Request):
    """Parse the AI-generated content into structured modules."""
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")
    content = m.get("content", "")
    modules = parse_mentorship_content(content)
    await db.mentorships.update_one({"id": mentorship_id}, {"$set": {"modules": modules}})
    return modules

def parse_mentorship_content(content: str) -> list:
    """Parse markdown content into structured modules."""
    import re
    modules = []
    # Split by module headers (## Modulo, **MODULO, etc.)
    module_pattern = r'(?:^|\n)(?:#{1,3}\s*|(?:\*\*))(?:M[oó]dulo\s*\d+|MODULO\s*\d+)[:\s\-–]*(.+?)(?:\*\*)?(?:\n|$)'
    parts = re.split(module_pattern, content, flags=re.IGNORECASE)

    if len(parts) < 2:
        # Try alternative split
        module_pattern2 = r'(?:^|\n)\d+\.\s*\*\*(.+?)\*\*'
        parts = re.split(module_pattern2, content)

    if len(parts) < 2:
        # If can't parse, create one big module
        modules.append({
            "id": str(uuid.uuid4()), "title": "Conteudo Completo",
            "objective": "", "order": 0,
            "lessons": [{"id": str(uuid.uuid4()), "title": "Conteudo", "content": content[:5000], "duration": "", "order": 0}],
            "exercises": [], "materials": []
        })
        return modules

    for i in range(1, len(parts), 2):
        title = parts[i].strip() if i < len(parts) else f"Modulo {len(modules)+1}"
        body = parts[i+1].strip() if i+1 < len(parts) else ""
        # Parse lessons from body
        lesson_pattern = r'(?:^|\n)\s*[-*]\s*(?:\*\*)?(?:Aula\s*\d+[:\s\-–]*)?(.+?)(?:\*\*)?(?:\n|$)'
        lesson_matches = re.findall(lesson_pattern, body, re.IGNORECASE)
        lessons = []
        if lesson_matches:
            for j, lt in enumerate(lesson_matches[:8]):
                lessons.append({"id": str(uuid.uuid4()), "title": lt.strip(), "content": "", "duration": "30min", "order": j})
        else:
            lessons.append({"id": str(uuid.uuid4()), "title": "Conteudo do modulo", "content": body[:2000], "duration": "60min", "order": 0})

        modules.append({
            "id": str(uuid.uuid4()), "title": title[:100],
            "objective": "", "order": len(modules),
            "lessons": lessons, "exercises": [], "materials": []
        })

    return modules[:12]  # Max 12 modules

# ─── Export ──────────────────────────────────────────────────────────────────
from fastapi.responses import Response
import markdown2

@router.get("/{mentorship_id}/export/pdf")
async def export_pdf(mentorship_id: str, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")

    html = build_export_html(m)
    try:
        from weasyprint import HTML
        pdf_bytes = HTML(string=html).write_pdf()
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{m["title"]}.pdf"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar PDF: {str(e)}")

@router.get("/{mentorship_id}/export/docx")
async def export_docx(mentorship_id: str, request: Request):
    user = await get_current_user(request)
    m = await db.mentorships.find_one({"id": mentorship_id, "user_id": user["_id"]})
    if not m:
        raise HTTPException(status_code=404, detail="Mentoria nao encontrada")

    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import io

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Title
        title_p = doc.add_heading(m.get("title", "Mentoria"), level=0)
        title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        if m.get("niche"):
            doc.add_paragraph(f"Nicho: {m['niche']} | Duracao: {m.get('duration_weeks', 8)} semanas", style='Subtitle')

        doc.add_paragraph("")

        # If has structured modules
        modules = m.get("modules", [])
        if modules:
            for mod in sorted(modules, key=lambda x: x.get("order", 0)):
                doc.add_heading(f"Modulo {mod['order']+1}: {mod['title']}", level=1)
                if mod.get("objective"):
                    doc.add_paragraph(f"Objetivo: {mod['objective']}")
                for lesson in sorted(mod.get("lessons", []), key=lambda x: x.get("order", 0)):
                    doc.add_heading(f"Aula: {lesson['title']}", level=2)
                    if lesson.get("content"):
                        doc.add_paragraph(lesson["content"])
                    if lesson.get("duration"):
                        doc.add_paragraph(f"Duracao: {lesson['duration']}", style='Intense Quote')
                if mod.get("exercises"):
                    doc.add_heading("Exercicios", level=2)
                    for ex in mod["exercises"]:
                        doc.add_paragraph(f"- {ex.get('title', '')}: {ex.get('description', '')}")
                doc.add_page_break()
        else:
            # Use raw content
            content = m.get("content", "")
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    p.add_run(line.strip("*")).bold = True
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return Response(
            content=buf.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{m["title"]}.docx"'}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {str(e)}")

def build_export_html(m: dict) -> str:
    """Build HTML profissional para export PDF — estilo eBook premium."""
    title = m.get("title", "Mentoria")
    modules = m.get("modules", [])
    content = m.get("content", "")
    niche = m.get("niche", "")
    audience = m.get("target_audience", "")
    weeks = m.get("duration_weeks", 8)
    from datetime import datetime
    date_str = datetime.now().strftime("%B de %Y")

    # TOC (table of contents)
    toc_items = ""
    if modules:
        for mod in sorted(modules, key=lambda x: x.get("order", 0)):
            toc_items += f'<li><span class="toc-num">{mod["order"]+1:02d}</span> <span class="toc-title">{mod["title"]}</span></li>'

    # Módulos formatados
    modules_html = ""
    if modules:
        for mod in sorted(modules, key=lambda x: x.get("order", 0)):
            modules_html += f'''
            <section class="module">
              <div class="mod-header">
                <span class="mod-num">{mod["order"]+1:02d}</span>
                <h2>{mod["title"]}</h2>
              </div>'''
            if mod.get("objective"):
                modules_html += f'<div class="objective"><strong>🎯 Objetivo</strong><p>{mod["objective"]}</p></div>'
            if mod.get("lessons"):
                modules_html += '<div class="lessons-list">'
                for lesson in sorted(mod.get("lessons", []), key=lambda x: x.get("order", 0)):
                    modules_html += f'<article class="lesson"><h3>{lesson["title"]}</h3>'
                    if lesson.get("duration"):
                        modules_html += f'<span class="duration">⏱ {lesson["duration"]}</span>'
                    if lesson.get("content"):
                        modules_html += f'<p>{lesson["content"]}</p>'
                    modules_html += '</article>'
                modules_html += '</div>'
            if mod.get("exercises"):
                modules_html += '<div class="exercises"><h3>📝 Exercícios práticos</h3>'
                for ex in mod["exercises"]:
                    modules_html += f'<div class="exercise"><strong>{ex.get("title","")}</strong><p>{ex.get("description","")}</p></div>'
                modules_html += '</div>'
            modules_html += '</section>'
    else:
        modules_html = markdown2.markdown(content, extras=["tables", "fenced-code-blocks"])

    html = f"""<!DOCTYPE html><html lang="pt-BR"><head><meta charset="utf-8">
    <title>{title}</title>
    <style>
      @page {{
        size: A4;
        margin: 2.2cm 1.8cm 2.5cm 1.8cm;
        @bottom-center {{
          content: counter(page) " de " counter(pages);
          font-family: 'Helvetica', sans-serif;
          font-size: 9pt;
          color: #888;
        }}
        @bottom-right {{
          content: "Kaelum.AI · Mentoria Premium";
          font-family: 'Helvetica', sans-serif;
          font-size: 8pt;
          color: #aaa;
        }}
      }}
      @page :first {{
        margin: 0;
        @bottom-center {{ content: none; }}
        @bottom-right {{ content: none; }}
      }}

      * {{ box-sizing: border-box; }}
      body {{
        font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.65;
        color: #2c2c2c;
        margin: 0;
      }}

      /* ─── CAPA ─── */
      .cover {{
        page-break-after: always;
        height: 29.7cm;
        padding: 4cm 2.5cm 3cm 2.5cm;
        background: linear-gradient(135deg, #0b1220 0%, #1a2744 50%, #2d1b3d 100%);
        color: #fff;
        position: relative;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
      }}
      .cover::before {{
        content: '';
        position: absolute;
        top: 0; right: 0;
        width: 8cm; height: 8cm;
        background: radial-gradient(circle, rgba(255,214,0,0.35) 0%, transparent 70%);
      }}
      .cover-brand {{
        position: relative;
        font-size: 11pt;
        letter-spacing: 6px;
        color: #FFD600;
        font-weight: 700;
        text-transform: uppercase;
      }}
      .cover-brand::before {{
        content: "■";
        margin-right: 10px;
        color: #FFD600;
      }}
      .cover-content {{ position: relative; }}
      .cover-label {{
        font-size: 10pt;
        letter-spacing: 4px;
        color: #FFD600;
        text-transform: uppercase;
        margin-bottom: 20px;
      }}
      .cover h1 {{
        font-size: 42pt;
        font-weight: 800;
        line-height: 1.1;
        margin: 0 0 30px 0;
        letter-spacing: -1.5px;
        color: #fff;
      }}
      .cover-subtitle {{
        font-size: 14pt;
        color: rgba(255,255,255,0.8);
        line-height: 1.5;
        margin-bottom: 40px;
        max-width: 14cm;
      }}
      .cover-meta {{
        position: relative;
        border-top: 1px solid rgba(255,214,0,0.3);
        padding-top: 25px;
        display: flex;
        justify-content: space-between;
        font-size: 10pt;
        color: rgba(255,255,255,0.7);
      }}
      .cover-meta div strong {{
        display: block;
        color: #FFD600;
        font-size: 9pt;
        letter-spacing: 2px;
        margin-bottom: 4px;
        text-transform: uppercase;
      }}

      /* ─── TOC ─── */
      .toc {{
        page-break-after: always;
        padding-top: 0.5cm;
      }}
      .toc-title {{
        font-size: 9pt;
        letter-spacing: 4px;
        color: #888;
        text-transform: uppercase;
        margin-bottom: 8px;
      }}
      .toc h2 {{
        font-size: 32pt;
        font-weight: 800;
        color: #0b1220;
        margin: 0 0 40px 0;
        letter-spacing: -1px;
        border: none;
        padding: 0;
      }}
      .toc ol {{ list-style: none; padding: 0; margin: 0; }}
      .toc li {{
        display: flex;
        align-items: baseline;
        padding: 14px 0;
        border-bottom: 1px solid #eee;
      }}
      .toc .toc-num {{
        font-weight: 800;
        color: #FFD600;
        font-size: 13pt;
        min-width: 2cm;
        letter-spacing: 1px;
      }}
      .toc .toc-title {{
        font-size: 12pt;
        color: #2c2c2c;
        font-weight: 500;
        letter-spacing: normal;
        text-transform: none;
      }}

      /* ─── MÓDULOS ─── */
      .module {{
        page-break-inside: auto;
        page-break-before: always;
        padding-top: 0.3cm;
      }}
      .mod-header {{
        display: flex;
        align-items: center;
        gap: 20px;
        margin-bottom: 30px;
        padding-bottom: 20px;
        border-bottom: 3px solid #FFD600;
      }}
      .mod-num {{
        background: #0b1220;
        color: #FFD600;
        font-weight: 800;
        font-size: 14pt;
        padding: 8px 14px;
        letter-spacing: 1px;
      }}
      .module h2 {{
        font-size: 22pt;
        font-weight: 800;
        color: #0b1220;
        margin: 0;
        line-height: 1.2;
        letter-spacing: -0.5px;
      }}

      .objective {{
        background: #fffbe6;
        border-left: 4px solid #FFD600;
        padding: 15px 20px;
        margin-bottom: 28px;
      }}
      .objective strong {{
        font-size: 10pt;
        letter-spacing: 2px;
        color: #8a6d00;
        text-transform: uppercase;
      }}
      .objective p {{ margin: 6px 0 0 0; font-size: 11pt; color: #2c2c2c; }}

      /* ─── LIÇÕES ─── */
      .lessons-list {{ margin-bottom: 30px; }}
      .lesson {{
        page-break-inside: avoid;
        margin-bottom: 22px;
        padding: 20px 24px;
        background: #fafafa;
        border-left: 4px solid #0b1220;
        position: relative;
      }}
      .lesson h3 {{
        font-size: 14pt;
        font-weight: 700;
        color: #0b1220;
        margin: 0 0 8px 0;
      }}
      .lesson .duration {{
        display: inline-block;
        background: #0b1220;
        color: #FFD600;
        font-size: 8pt;
        letter-spacing: 1px;
        padding: 3px 10px;
        margin-bottom: 10px;
        font-weight: 700;
      }}
      .lesson p {{
        margin: 0;
        font-size: 10.5pt;
        color: #444;
        line-height: 1.7;
      }}

      /* ─── EXERCÍCIOS ─── */
      .exercises {{
        page-break-inside: avoid;
        background: linear-gradient(135deg, #fff8d6 0%, #fff 100%);
        border: 2px dashed #FFD600;
        padding: 25px;
        margin-top: 30px;
      }}
      .exercises h3 {{
        font-size: 13pt;
        color: #8a6d00;
        margin: 0 0 15px 0;
        font-weight: 800;
      }}
      .exercise {{
        background: #fff;
        padding: 14px 18px;
        margin-bottom: 12px;
        border-left: 3px solid #FFD600;
      }}
      .exercise:last-child {{ margin-bottom: 0; }}
      .exercise strong {{
        display: block;
        color: #0b1220;
        margin-bottom: 4px;
        font-size: 11pt;
      }}
      .exercise p {{
        margin: 0;
        font-size: 10pt;
        color: #555;
      }}

      /* Raw markdown fallback */
      h2 {{ color: #0b1220; font-size: 18pt; border-bottom: 2px solid #FFD600; padding-bottom: 8px; margin-top: 40px; }}
      h3 {{ color: #333; font-size: 13pt; margin-top: 20px; }}
      p {{ margin: 10px 0; }}
      ul, ol {{ padding-left: 24px; }}
      strong {{ color: #0b1220; }}
      code {{ background: #f0f0f0; padding: 2px 6px; border-radius: 3px; font-size: 9pt; }}
      blockquote {{
        border-left: 4px solid #FFD600;
        padding: 8px 16px;
        color: #555;
        font-style: italic;
        background: #fffbe6;
        margin: 12px 0;
      }}
    </style></head><body>

    <!-- CAPA -->
    <div class="cover">
      <div class="cover-brand">Kaelum.AI</div>
      <div class="cover-content">
        <div class="cover-label">Mentoria Premium · {weeks} semanas</div>
        <h1>{title}</h1>
        <p class="cover-subtitle">{(audience or niche or 'Programa estruturado em módulos progressivos para transformação prática.')[:180]}</p>
      </div>
      <div class="cover-meta">
        <div><strong>Nicho</strong>{niche or '—'}</div>
        <div><strong>Duração</strong>{weeks} semanas</div>
        <div><strong>Edição</strong>{date_str}</div>
      </div>
    </div>

    <!-- SUMÁRIO -->
    <div class="toc">
      <div class="toc-title">Índice</div>
      <h2>Conteúdo Programático</h2>
      <ol>
        {toc_items}
      </ol>
    </div>

    <!-- MÓDULOS -->
    {modules_html}

    </body></html>"""
    return html
